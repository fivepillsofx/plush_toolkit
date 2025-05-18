import os
from io import BytesIO

BASE = os.path.dirname(os.path.abspath(__file__))
PAGES = os.path.join(BASE, "pages")
os.makedirs(PAGES, exist_ok=True)

# 1) plush_utils.py
with open(os.path.join(BASE, "plush_utils.py"), "w") as f:
    f.write("""\
import re, random, textstat, nltk
from nltk.data import find
from nltk.tokenize.punkt import PunktSentenceTokenizer
from nltk.tokenize import wordpunct_tokenize
from nltk.corpus import stopwords
from striprtf.striprtf import rtf_to_text
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from datetime import datetime
from io import BytesIO
from collections import Counter

def ensure_nltk_data():
    try: find("tokenizers/punkt")
    except: nltk.download("punkt", quiet=True)
    try: find("corpora/stopwords")
    except: nltk.download("stopwords", quiet=True)
ensure_nltk_data()

FILLER_WORDS = ["just","really","very","that","actually","like","maybe","somewhat","perhaps","quite"]
CLICHES = [ 
    "needle in a haystack","cold sweat","chill ran down","time stood still","dead silence",
    "at the end of the day","low-hanging fruit","the calm before the storm","head over heels",
    "in the nick of time","plenty of fish in the sea","easy as pie","scared stiff",
    "raining cats and dogs","think outside the box","every cloud has a silver lining",
    "pushing up daisies","barking up the wrong tree","blood ran cold","fit as a fiddle"
]
STYLE_PRESETS = {
    "None":      {"emphasis":"","note":""},
    "Gritty":    {"emphasis":"Cliché detection, passive voice, long sentences","note":""},
    "Snappy":    {"emphasis":"Filler words, punchy structure","note":""},
    "Poetic":    {"emphasis":"Flow, sentence variety","note":""},
    "Technical": {"emphasis":"Passive voice, clarity","note":""},
    "Sparse":    {"emphasis":"Minimal filler, clarity","note":""}
}

COMMON_MALE   = ["James","John","Robert","Michael","William","David","Richard","Joseph","Thomas","Charles"]
COMMON_FEMALE = ["Mary","Patricia","Jennifer","Linda","Elizabeth","Barbara","Susan","Jessica","Sarah","Karen"]
RARE_MALE     = ["Ansel","Blaise","Caius","Dorian","Elwood","Fintan","Gideon","Ivo","Leander","Montague"]
RARE_FEMALE   = ["Aurelia","Briseis","Calista","Delphine","Elowen","Ferelith","Galatea","Isolde","Junia","Lysandra"]

templates = {
    "Three-Act Beat Sheet": "...",  # full text omitted for brevity
    "Scene & Chapter Planner": "..."
}

def load_docx(f): return "\\n".join(p.text for p in DocxDocument(f).paragraphs)

def generate_docx(title, content):
    doc = DocxDocument(); doc.add_heading(title,1)
    for L in content.split("\\n"): doc.add_paragraph(L)
    b=BytesIO(); doc.save(b); b.seek(0); return b.getvalue()

def generate_pdf(title, content):
    b=BytesIO(); c=canvas.Canvas(b,pagesize=letter)
    w,h=letter; tx=c.beginText(40,h-40)
    tx.setFont("Helvetica-Bold",14); tx.textLine(title)
    tx.setFont("Helvetica",12)
    for L in content.split("\\n"): tx.textLine(L)
    c.drawText(tx); c.showPage(); c.save(); b.seek(0); return b.getvalue()

def clean_text(txt):
    return " ".join(txt.replace("“","\"").replace("”","\"")
                   .replace("‘","'").replace("’","'")
                   .replace("--","—").split())

def detect_passive(txt):
    toks=PunktSentenceTokenizer().tokenize(txt)
    pat=re.compile(r'\\b(was|were).*?\\b\\w+ed\\b',re.I)
    return [(i+1,s.strip()) for i,s in enumerate(toks) if pat.search(s)]

def suggest(txt):
    out=[]; toks=PunktSentenceTokenizer().tokenize(txt)
    for i,s in enumerate(toks):
        issues=[]
        if len(wordpunct_tokenize(s))>30: issues.append("⚠️ Break it up")
        if sum(s.lower().count(w) for w in FILLER_WORDS)>2: issues.append("✂️ Cut filler")
        if re.search(r'\\b(was|were)\\b',s,re.I): issues.append("💡 Try active voice")
        if issues: out.append(f"Sentence {i+1}:\\n{s}\\n" + "\\n".join(issues))
    return "\\n\\n".join(out) if out else "✅ All good!"

def analyze_text(txt, style):
    rpt=[]
    if style!="None":
        rpt += [f"🎨 {style}", STYLE_PRESETS[style]["emphasis"], ""]
    rpt += [
      f"• Words: {len(txt.split())}",
      f"• Sents: {textstat.sentence_count(txt)}",
      f"• Avg len: {textstat.words_per_sentence(txt):.2f}",
      f"• Grade: {textstat.flesch_kincaid_grade(txt):.2f}", ""
    ]
    rpt.append("🔎 Fillers:")
    for w in FILLER_WORDS:
        c=txt.lower().split().count(w)
        if c: rpt.append(f" - {w}: {c}")
    rpt.append("\\n⚠️ Long sents:")
    for i,s in enumerate(PunktSentenceTokenizer().tokenize(txt)):
        if len(wordpunct_tokenize(s))>30: rpt.append(f"{i+1}: {s}")
    rpt.append("\\n🕵️ Passive:")
    pv=detect_passive(txt)
    rpt += [f"{n}: {s}" for n,s in pv] if pv else ["✅ None"]
    rpt.append("\\n🤖 Suggestions:\\n"+suggest(txt))
    return "\\n\\n".join(rpt)

def extract_dialogue(txt): return "\\n".join(re.findall(r'[“"]([^“”"]+)[”"]',txt))

def dialogue_by_character(txt):
    pat=r'"[^"]*?"\\s+(?:said|asked|replied)\\s+([A-Z][a-zA-Z]*)'
    names=re.findall(pat,txt)
    return "\\n".join(f"{n}: {c}" for n,c in Counter(names).items()) if names else "None"

def find_cliches(txt):
    found=[c for c in CLICHES if c in txt.lower()]
    return "\\n".join(found) if found else "None"

def export_full_report(txt, style):
    dt=datetime.now().strftime("%Y-%m-%d %H:%M")
    parts=[f"Full Report | {dt}", analyze_text(txt,style),
           "\\n=== Dialogue ===\\n"+dialogue_by_character(txt),
           "\\n=== Extracted ===\\n"+extract_dialogue(txt),
           "\\n=== Clichés ===\\n"+find_cliches(txt)]
    return "\\n\\n".join(parts)

def generate_names(gender, rarity, count):
    pool=[]
    if gender in ("Any","Male"): pool += (RARE_MALE if rarity=="Rare" else COMMON_MALE)
    if gender in ("Any","Female"): pool += (RARE_FEMALE if rarity=="Rare" else COMMON_FEMALE)
    return random.sample(pool, min(count,len(pool)))
""")

# 2) pages—just show the pattern, you can duplicate/change titles
page_tpl = """
import streamlit as st
from plush_utils import {imports}

st.set_page_config(page_title="{title}")
st.title("{emoji} {title}")

uploaded = st.file_uploader("Upload .txt, .docx, or .rtf", type=["txt","docx","rtf"])
if uploaded:
    ext=uploaded.name.split(".")[-1].lower()
    if ext=="txt": raw=uploaded.read().decode()
    elif ext=="docx": raw=load_docx(uploaded)
    else: raw=rtf_to_text(uploaded.read().decode())
else:
    raw=st.text_area("Or paste here", height=300)

if st.button("{btn_label}"):
    {action_block}
"""

pages = [
    ("01_Clean_Text.py",      "clean_text, load_docx, rtf_to_text, generate_docx, generate_pdf",         "Clean Text",    "Clean Text",      "out=clean_text(raw)\n    st.text_area('', out, height=300)\n    st.download_button('.txt',out,'cleaned.txt')"),
    ("02_Analyze_Text.py",    "analyze_text, STYLE_PRESETS, load_docx, rtf_to_text, generate_docx, generate_pdf", "Analyze Text",  "Analyze",         "report=analyze_text(raw, style)\n    st.text_area('', report, height=400)"),
    ("03_Extract_Dialogue.py","extract_dialogue, load_docx, rtf_to_text",                                "Extract Dialogue","Extract Dialogue","d=extract_dialogue(raw)\n    st.text_area('', d, height=300)"),
    ("04_Dialogue_by_Character.py","dialogue_by_character, load_docx, rtf_to_text",                "Dialogue by Character","Show Dialogue","r=dialogue_by_character(raw)\n    st.text_area('', r, height=300)"),
    ("05_Cliche_Buster.py",   "find_cliches, load_docx, rtf_to_text",                              "Cliché Buster", "Find Clichés",    "c=find_cliches(raw)\n    st.text_area('', c, height=300)"),
    ("06_Full_Report.py",     "export_full_report, STYLE_PRESETS, load_docx, rtf_to_text, generate_docx, generate_pdf", "Full Report","Generate Full Report","fr=export_full_report(raw,style)\n    st.text_area('', fr, height=400)"),
    ("07_Character_Name_Generator.py","generate_names",                                          "Character Name Generator","Generate","names=generate_names(gender,rarity,count)\n    st.write('\\n'.join(names))"),
    ("08_Templates.py",       "templates",                                                    "Templates",      "Show Template",  "tpl=templates[choice]\n    st.text_area('', tpl, height=400)")
]

for filename, imports, title, btn, action in pages:
    content = page_tpl.format(imports=imports, title=title, emoji=btn.split()[0], btn_label=btn, action_block=action)
    with open(os.path.join(PAGES, filename), "w") as f:
        f.write(content)

print("All files generated. You now have plush_utils.py and pages/01...08 ready.")