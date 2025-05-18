import streamlit as st
import os
import re
import random
from io import BytesIO
from collections import Counter
from datetime import datetime
import nltk
from nltk.data import find
from nltk.tokenize.punkt import PunktSentenceTokenizer
from nltk.tokenize import wordpunct_tokenize
from nltk.corpus import stopwords
import textstat
from docx import Document as DocxDocument
from striprtf.striprtf import rtf_to_text
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# ── Ensure NLTK Data ───────────────────────────────────────
def ensure_nltk_data():
    try:
        find("tokenizers/punkt")
    except:
        nltk.download("punkt", quiet=True)
    try:
        find("corpora/stopwords")
    except:
        nltk.download("stopwords", quiet=True)

ensure_nltk_data()

# ── Constants ───────────────────────────────────────────────
FILLER_WORDS = ["just", "really", "very", "that", "actually", "like", "maybe", "somewhat", "perhaps", "quite"]

CLICHES = [
    "needle in a haystack", "cold sweat", "chill ran down", "time stood still", "dead silence",
    "at the end of the day", "low-hanging fruit", "the calm before the storm", "head over heels",
    "in the nick of time", "plenty of fish in the sea", "easy as pie", "scared stiff",
    "raining cats and dogs", "think outside the box", "every cloud has a silver lining",
    "pushing up daisies", "barking up the wrong tree", "blood ran cold", "fit as a fiddle"
]

STYLE_PRESETS = {
    "None":      {"emphasis": "", "note": ""},
    "Gritty":    {"emphasis": "Cliché detection, passive voice, long sentences", "note": "Highlights realism."},
    "Snappy":    {"emphasis": "Filler words, punchy structure", "note": "Focuses on rhythm."},
    "Poetic":    {"emphasis": "Flow, sentence variety", "note": "Flags broken rhythm."},
    "Technical": {"emphasis": "Passive voice, clarity", "note": "Suited for nonfiction."},
    "Sparse":    {"emphasis": "Minimal filler, clarity", "note": "Zero fluff."}
}

COMMON_MALE   = ["James","John","Robert","Michael","William","David","Richard","Joseph","Thomas","Charles"]
COMMON_FEMALE = ["Mary","Patricia","Jennifer","Linda","Elizabeth","Barbara","Susan","Jessica","Sarah","Karen"]
RARE_MALE     = ["Ansel","Blaise","Caius","Dorian","Elwood","Fintan","Gideon","Ivo","Leander","Montague"]
RARE_FEMALE   = ["Aurelia","Briseis","Calista","Delphine","Elowen","Ferelith","Galatea","Isolde","Junia","Lysandra"]

templates = {
    "Three-Act Beat Sheet": """# Three-Act Beat Sheet

## Act I: Setup
- **Opening Image**:  
- **Inciting Incident**:  
- **Key Event**:  

## Act II: Confrontation
- **Pinch Point 1**:  
- **Midpoint**:  
- **Pinch Point 2**:  

## Act III: Resolution
- **Climax**:  
- **Denouement**:  
""",
    "Scene & Chapter Planner": """# Scene & Chapter Planner

## Scene / Chapter Title

### Goal  
What does your protagonist want?

### Conflict  
What stands in their way?

### Stakes  
What's at risk?

### Outcome  
How does this scene change the story?

## Notes  
- Setting  
- POV Character  
- Key Dialogue Beats  
- Mood / Tone  
"""
}

# ── Helper Functions ────────────────────────────────────────
def load_docx(f):
    return "\n".join(p.text for p in DocxDocument(f).paragraphs)

def generate_docx(title, content):
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for line in content.split("\n"):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def generate_pdf(title, content):
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    w, h = letter
    text = c.beginText(40, h - 40)
    text.setFont("Helvetica-Bold", 14)
    text.textLine(title)
    text.setFont("Helvetica", 12)
    for line in content.split("\n"):
        text.textLine(line)
    c.drawText(text)
    c.showPage()
    c.save()
    buf.seek(0)
    return buf.getvalue()

def clean_text(txt):
    return " ".join(
        txt.replace("“", "\"")
           .replace("”", "\"")
           .replace("‘", "'")
           .replace("’", "'")
           .replace("--", "—")
           .split()
    )

def detect_passive(txt):
    toks = PunktSentenceTokenizer().tokenize(txt)
    pat = re.compile(r'\b(was|were|is being|are being|has been|have been|had been)\b\s+\w+ed\b', re.IGNORECASE)
    return [(i + 1, s.strip()) for i, s in enumerate(toks) if pat.search(s)]

def suggest_improvements(txt):
    out = []
    toks = PunktSentenceTokenizer().tokenize(txt)
    for i, s in enumerate(toks):
        issues = []
        if len(wordpunct_tokenize(s)) > 30:
            issues.append("⚠️ Break up this sentence.")
        if sum(s.lower().count(w) for w in FILLER_WORDS) > 2:
            issues.append("✂️ Cut filler words.")
        if re.search(r'\b(was|were)\b', s, re.IGNORECASE):
            issues.append("💡 Try active voice.")
        if issues:
            out.append(f"Sentence {i+1}:\n{s}\n" + "\n".join(issues))
    return "\n\n".join(out) if out else "✅ All good!"

def analyze_text(txt, style):
    rpt = []
    if style != "None":
        rpt += [f"🎨 Style: {style}", STYLE_PRESETS[style]["emphasis"], ""]
    rpt += [
        f"• Words: {len(txt.split())}",
        f"• Sentences: {textstat.sentence_count(txt)}",
        f"• Avg length: {textstat.words_per_sentence(txt):.2f}",
        f"• Grade level: {textstat.flesch_kincaid_grade(txt):.2f}",
        ""
    ]
    rpt.append("🔎 Filler Words:")
    for w in FILLER_WORDS:
        c = txt.lower().split().count(w)
        if c:
            rpt.append(f" - {w}: {c}")
    rpt.append("\n⚠️ Long Sentences:")
    for i, s in enumerate(PunktSentenceTokenizer().tokenize(txt)):
        if len(wordpunct_tokenize(s)) > 30:
            rpt.append(f"{i+1}: {s}")
    rpt.append("\n🕵️ Passive Voice:")
    pv = detect_passive(txt)
    rpt += [f"{n}: {s}" for n, s in pv] if pv else ["✅ None"]
    rpt.append("\n🤖 Smart Suggestions:\n" + suggest_improvements(txt))
    return "\n\n".join(rpt)

def extract_dialogue(txt):
    return "\n".join(re.findall(r'[“"]([^“”"]+)[”"]', txt))

def dialogue_by_character(txt):
    pat = r'"[^"]+?"\s+(?:said|asked|replied|whispered|shouted|cried|muttered|yelled|snapped|called)\s+([A-Z][a-zA-Z]*)'
    names = re.findall(pat, txt)
    if not names:
        return "None found."
    cnt = Counter(names)
    return "\n".join(f"{n}: {c}" for n, c in cnt.items())

def find_cliches(txt):
    found = [c for c in CLICHES if c in txt.lower()]
    return "\n".join(found) if found else "None"

def export_full_report(txt, style):
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    parts = [
        f"Full Report | Generated {now}",
        analyze_text(txt, style),
        "\n=== Dialogue by Character ===\n" + dialogue_by_character(txt),
        "\n=== Extracted Dialogue ===\n" + extract_dialogue(txt),
        "\n=== Clichés ===\n" + find_cliches(txt)
    ]
    return "\n\n".join(parts)

def generate_names(gender, rarity, count):
    pool = []
    if gender in ("Any", "Male"):
        pool += (RARE_MALE if rarity == "Rare" else COMMON_MALE)
    if gender in ("Any", "Female"):
        pool += (RARE_FEMALE if rarity == "Rare" else COMMON_FEMALE)
    return random.sample(pool, min(count, len(pool)))

# ── Main App ────────────────────────────────────────────────
def main():
    st.set_page_config(page_title="🧠 Plush Toolkit", layout="wide")

    tools = [
        "Home",
        "Clean Text",
        "Analyze Text",
        "Extract Dialogue",
        "Dialogue by Character",
        "Cliché Buster",
        "Full Report",
        "Character Name Generator",
        "Templates"
    ]
    choice = st.sidebar.radio("Select a tool", tools)
    c = st.container()

    if choice == "Home":
        c.title("📚 Plush: The Writer Toolkit")
        c.write("Welcome! Use the sidebar to pick a tool.")

    elif choice == "Clean Text":
        c.title("🧼 Clean Text")
        f = c.file_uploader("Upload `.txt`, `.docx`, or `.rtf`", type=["txt", "docx", "rtf"])
        if f:
            ext = f.name.split(".")[-1].lower()
            raw = (
                f.read().decode("utf-8")
                if ext == "txt"
                else load_docx(f)
                if ext == "docx"
                else rtf_to_text(f.read().decode("utf-8"))
            )
        else:
            raw = c.text_area("Or paste your text here:", height=300)
        if c.button("🧼 Clean"):
            out = clean_text(raw)
            c.subheader("✅ Cleaned Text")
            c.text_area("", out, height=300)
            c.download_button("Download .txt", out, "cleaned.txt")
            c.download_button("Download .docx", generate_docx("Cleaned Text", out), "cleaned.docx")
            c.download_button("Download .pdf", generate_pdf("Cleaned Text", out), "cleaned.pdf")

    elif choice == "Analyze Text":
        c.title("🔍 Analyze Text")
        f = c.file_uploader("Upload `.txt`, `.docx`, or `.rtf`", type=["txt", "docx", "rtf"])
        if f:
            ext = f.name.split(".")[-1].lower()
            raw = (
                f.read().decode("utf-8")
                if ext == "txt"
                else load_docx(f)
                if ext == "docx"
                else rtf_to_text(f.read().decode("utf-8"))
            )
        else:
            raw = c.text_area("Or paste your text here:", height=300)
        style = c.selectbox("Style Preset", list(STYLE_PRESETS.keys()))
        if c.button("🔍 Analyze"):
            rpt = analyze_text(raw, style)
            c.subheader("📊 Analysis Report")
            c.text_area("", rpt, height=400)
            c.download_button("Download .txt", rpt, "report.txt")
            c.download_button("Download .docx", generate_docx("Analysis Report", rpt), "analysis.docx")
            c.download_button("Download .pdf", generate_pdf("Analysis Report", rpt), "analysis.pdf")

    elif choice == "Extract Dialogue":
        c.title("🗣 Extract Dialogue")
        f = c.file_uploader("Upload `.txt`, `.docx`, or `.rtf`", type=["txt", "docx", "rtf"])
        if f:
            ext = f.name.split(".")[-1].lower()
            raw = (
                f.read().decode("utf-8")
                if ext == "txt"
                else load_docx(f)
                if ext == "docx"
                else rtf_to_text(f.read().decode("utf-8"))
            )
        else:
            raw = c.text_area("Or paste your text here:", height=300)
        if c.button("🗣 Extract"):
            dlg = extract_dialogue(raw)
            c.subheader("🗣 Extracted Dialogue")
            c.text_area("", dlg, height=300)
            c.download_button("Download .txt", dlg, "dialogue.txt")

    elif choice == "Dialogue by Character":
        c.title("🧍 Dialogue by Character")
        f = c.file_uploader("Upload `.txt`, `.docx`, or `.rtf`", type=["txt", "docx", "rtf"])
        if f:
            ext = f.name.split(".")[-1].lower()
            raw = (
                f.read().decode("utf-8")
                if ext == "txt"
                else load_docx(f)
                if ext == "docx"
                else rtf_to_text(f.read().decode("utf-8"))
            )
        else:
            raw = c.text_area("Or paste your text here:", height=300)
        if c.button("🧍 Show"):
            rep = dialogue_by_character(raw)
            c.subheader("🧍 Dialogue by Character")
            c.text_area("", rep, height=300)
            c.download_button("Download .txt", rep, "by_character.txt")

    elif choice == "Cliché Buster":
        c.title("💣 Cliché Buster")
        f = c.file_uploader("Upload `.txt`, `.docx`, or `.rtf`", type=["txt", "docx", "rtf"])
        if f:
            ext = f.name.split(".")[-1].lower()
            raw = (
                f.read().decode("utf-8")
                if ext == "txt"
                else load_docx(f)
                if ext == "docx"
                else rtf_to_text(f.read().decode("utf-8"))
            )
        else:
            raw = c.text_area("Or paste your text here:", height=300)
        if c.button("💣 Bust"):
            rep = find_cliches(raw)
            c.subheader("💣 Clichés Found")
            c.text_area("", rep, height=300)
            c.download_button("Download .txt", rep, "cliches.txt")

    elif choice == "Full Report":
        c.title("📦 Full Report")
        f = c.file_uploader("Upload `.txt`, `.docx`, or `.rtf`", type=["txt", "docx", "rtf"])
        if f:
            ext = f.name.split(".")[-1].lower()
            raw = (
                f.read().decode("utf-8")
                if ext == "txt"
                else load_docx(f)
                if ext == "docx"
                else rtf_to_text(f.read().decode("utf-8"))
            )
        else:
            raw = c.text_area("Or paste your text here:", height=300)
        style = c.selectbox("Style Preset", list(STYLE_PRESETS.keys()))
        if c.button("Generate Report"):
            fr = export_full_report(raw, style)
            c.subheader("📦 Full Report")
            c.text_area("", fr, height=400)
            c.download_button("Download .txt", fr, "full_report.txt")
            c.download_button("Download .docx", generate_docx("Full Report", fr), "full_report.docx")
            c.download_button("Download .pdf", generate_pdf("Full Report", fr), "full_report.pdf")

    elif choice == "Character Name Generator":
        c.title("🎲 Character Name Generator")
        gender = c.selectbox("Gender", ["Any", "Male", "Female"])
        rarity = c.selectbox("Rarity", ["Common", "Rare"])
        count = c.number_input("How many names?", 1, 10, 5)
        if c.button("🎲 Generate"):
            names = generate_names(gender, rarity, count)
            c.subheader("🎲 Generated Names")
            for n in names:
                c.write(f"- {n}")

    elif choice == "Templates":
        c.title("📑 Templates Library")
        tpl_choice = c.selectbox("Pick a template", list(templates.keys()))
        if c.button("Show Template"):
            t = templates[tpl_choice]
            c.subheader(f"📑 {tpl_choice}")
            c.text_area("", t, height=400)
            c.download_button("Download .txt", t, f"{tpl_choice}.txt")
            c.download_button("Download .docx", generate_docx(tpl_choice, t), f"{tpl_choice}.docx")
            c.download_button("Download .pdf", generate_pdf(tpl_choice, t), f"{tpl_choice}.pdf")

if __name__ == "__main__":
    main()
