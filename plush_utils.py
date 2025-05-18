import re, random, textstat, nltk
from nltk.data import find
from nltk.tokenize.punkt import PunktSentenceTokenizer
from nltk.tokenize import wordpunct_tokenize
from nltk.corpus import stopwords
from striprtf.striprtf import rtf_to_text
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from datetime import datetime
from io import BytesIO
from collections import Counter

def ensure_nltk_data():
    try: find("tokenizers/punkt")
    except: nltk.download("punkt", quiet=True)
    try: find("corpora/stopwords")
    except: nltk.download("stopwords", quiet=True)
ensure_nltk_data()

FILLER_WORDS = ["just","really","very","that","actually","like","maybe","somewhat","perhaps","quite"]
CLICHES = [ 
    "needle in a haystack","cold sweat","chill ran down","time stood still","dead silence",
    "at the end of the day","low-hanging fruit","the calm before the storm","head over heels",
    "in the nick of time","plenty of fish in the sea","easy as pie","scared stiff",
    "raining cats and dogs","think outside the box","every cloud has a silver lining",
    "pushing up daisies","barking up the wrong tree","blood ran cold","fit as a fiddle"
]
STYLE_PRESETS = {
    "None":      {"emphasis":"","note":""},
    "Gritty":    {"emphasis":"Cliché detection, passive voice, long sentences","note":""},
    "Snappy":    {"emphasis":"Filler words, punchy structure","note":""},
    "Poetic":    {"emphasis":"Flow, sentence variety","note":""},
    "Technical": {"emphasis":"Passive voice, clarity","note":""},
    "Sparse":    {"emphasis":"Minimal filler, clarity","note":""}
}

COMMON_MALE   = ["James","John","Robert","Michael","William","David","Richard","Joseph","Thomas","Charles"]
COMMON_FEMALE = ["Mary","Patricia","Jennifer","Linda","Elizabeth","Barbara","Susan","Jessica","Sarah","Karen"]
RARE_MALE     = ["Ansel","Blaise","Caius","Dorian","Elwood","Fintan","Gideon","Ivo","Leander","Montague"]
RARE_FEMALE   = ["Aurelia","Briseis","Calista","Delphine","Elowen","Ferelith","Galatea","Isolde","Junia","Lysandra"]

templates = {
    "Three-Act Beat Sheet": "...",  # full text omitted for brevity
    "Scene & Chapter Planner": "..."
}

def load_docx(f): return "\n".join(p.text for p in DocxDocument(f).paragraphs)

def generate_docx(title, content):
    doc = DocxDocument(); doc.add_heading(title,1)
    for L in content.split("\n"): doc.add_paragraph(L)
    b=BytesIO(); doc.save(b); b.seek(0); return b.getvalue()

def generate_pdf(title, content):
    b=BytesIO(); c=canvas.Canvas(b,pagesize=letter)
    w,h=letter; tx=c.beginText(40,h-40)
    tx.setFont("Helvetica-Bold",14); tx.textLine(title)
    tx.setFont("Helvetica",12)
    for L in content.split("\n"): tx.textLine(L)
    c.drawText(tx); c.showPage(); c.save(); b.seek(0); return b.getvalue()

def clean_text(txt):
    return " ".join(txt.replace("“",""").replace("”",""")
                   .replace("‘","'").replace("’","'")
                   .replace("--","—").split())

def detect_passive(txt):
    toks=PunktSentenceTokenizer().tokenize(txt)
    pat=re.compile(r'\b(was|were).*?\b\w+ed\b',re.I)
    return [(i+1,s.strip()) for i,s in enumerate(toks) if pat.search(s)]

def suggest(txt):
    out=[]; toks=PunktSentenceTokenizer().tokenize(txt)
    for i,s in enumerate(toks):
        issues=[]
        if len(wordpunct_tokenize(s))>30: issues.append("⚠️ Break it up")
        if sum(s.lower().count(w) for w in FILLER_WORDS)>2: issues.append("✂️ Cut filler")
        if re.search(r'\b(was|were)\b',s,re.I): issues.append("💡 Try active voice")
        if issues: out.append(f"Sentence {i+1}:\n{s}\n" + "\n".join(issues))
    return "\n\n".join(out) if out else "✅ All good!"

def analyze_text(txt, style):
    rpt=[]
    if style!="None":
        rpt += [f"🎨 {style}", STYLE_PRESETS[style]["emphasis"], ""]
    rpt += [
      f"• Words: {len(txt.split())}",
      f"• Sents: {textstat.sentence_count(txt)}",
      f"• Avg len: {textstat.words_per_sentence(txt):.2f}",
      f"• Grade: {textstat.flesch_kincaid_grade(txt):.2f}", ""
    ]
    rpt.append("🔎 Fillers:")
    for w in FILLER_WORDS:
        c=txt.lower().split().count(w)
        if c: rpt.append(f" - {w}: {c}")
    rpt.append("\n⚠️ Long sents:")
    for i,s in enumerate(PunktSentenceTokenizer().tokenize(txt)):
        if len(wordpunct_tokenize(s))>30: rpt.append(f"{i+1}: {s}")
    rpt.append("\n🕵️ Passive:")
    pv=detect_passive(txt)
    rpt += [f"{n}: {s}" for n,s in pv] if pv else ["✅ None"]
    rpt.append("\n🤖 Suggestions:\n"+suggest(txt))
    return "\n\n".join(rpt)

def extract_dialogue(txt): return "\n".join(re.findall(r'[“"]([^“”"]+)[”"]',txt))

def dialogue_by_character(txt):
    pat=r'"[^"]*?"\s+(?:said|asked|replied)\s+([A-Z][a-zA-Z]*)'
    names=re.findall(pat,txt)
    return "\n".join(f"{n}: {c}" for n,c in Counter(names).items()) if names else "None"

def find_cliches(txt):
    found=[c for c in CLICHES if c in txt.lower()]
    return "\n".join(found) if found else "None"

def export_full_report(txt, style):
    dt=datetime.now().strftime("%Y-%m-%d %H:%M")
    parts=[f"Full Report | {dt}", analyze_text(txt,style),
           "\n=== Dialogue ===\n"+dialogue_by_character(txt),
           "\n=== Extracted ===\n"+extract_dialogue(txt),
           "\n=== Clichés ===\n"+find_cliches(txt)]
    return "\n\n".join(parts)

def generate_names(gender, rarity, count):
    pool=[]
    if gender in ("Any","Male"): pool += (RARE_MALE if rarity=="Rare" else COMMON_MALE)
    if gender in ("Any","Female"): pool += (RARE_FEMALE if rarity=="Rare" else COMMON_FEMALE)
    return random.sample(pool, min(count,len(pool)))
